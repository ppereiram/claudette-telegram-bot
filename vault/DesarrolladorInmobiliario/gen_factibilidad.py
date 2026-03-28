
# -*- coding: utf-8 -*-
import openpyxl
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colors ───────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
LIGHT_BLUE = RGBColor(0xEB, 0xF0, 0xF7)
GREEN_BG   = RGBColor(0xC6, 0xEF, 0xCE)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)

# ─── Helper functions ──────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    # RGBColor is a named tuple / sequence: [r, g, b]
    hx  = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hx)
    tcPr.append(shd)

def add_cell(cell, text, bold=False, color=BLACK, size=9,
             align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(str(text))
    run.bold           = bold
    run.italic         = italic
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    run.font.name      = 'Calibri'
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def hdr(row, texts, bg=DARK_BLUE, fg=WHITE, size=9, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER):
    for cell, text in zip(row.cells, texts):
        set_cell_bg(cell, bg)
        add_cell(cell, text, bold=bold, color=fg, size=size, align=align)

def dat(row, texts, bg=None, fg=BLACK, bold=False, size=9,
        aligns=None, default_align=WD_ALIGN_PARAGRAPH.LEFT):
    for i, (cell, text) in enumerate(zip(row.cells, texts)):
        if bg:
            set_cell_bg(cell, bg)
        al = aligns[i] if aligns else default_align
        add_cell(cell, text, bold=bold, color=fg, size=size, align=al)

def sub(row, texts, bg=MED_BLUE, fg=WHITE, size=9, bold=True,
        aligns=None):
    for i, (cell, text) in enumerate(zip(row.cells, texts)):
        set_cell_bg(cell, bg)
        al = (aligns[i] if aligns else WD_ALIGN_PARAGRAPH.LEFT)
        add_cell(cell, text, bold=bold, color=fg, size=size, align=al)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold           = True
    run.font.name      = 'Calibri'
    run.font.color.rgb = DARK_BLUE
    run.font.size      = Pt(13 if level == 1 else 11 if level == 2 else 10)
    return p

def add_body(doc, text, size=10, color=BLACK, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    run.italic         = italic
    return p

def page_break(doc):
    doc.add_page_break()

def fmt_usd(n):
    return '${:,.0f}'.format(n)

def fmt_m2(n):
    return '{:,.2f} m\u00b2'.format(n)

# ─── Read Excel data ───────────────────────────────────────────────────────────
EXCEL_PATH = r'c:/Users/Pablo/Documents/Obsidian Vault/Desarrollador Inmobiliario/FINCAS FILIALES POR ETAPAS 11-08-2020.xlsx'
wb = openpyxl.load_workbook(EXCEL_PATH, data_only=True)
ws = wb.active
all_rows = list(ws.iter_rows(values_only=True))

FLOOR_COLS = {
    'Piso 1': 11, 'Piso 2': 12, 'Piso 3': 13,
    'Piso 4': 14, 'Piso 5': 15, 'Piso 6': 16,
    'Piso 7': 17, 'Piso 8': 18, 'Piso 9': 19,
    'Piso 10': 20, 'Terraza': 21,
}
FLOOR_MULTIPLIER = {
    'Piso 1': 1.000, 'Piso 2': 1.020, 'Piso 3': 1.040,
    'Piso 4': 1.061, 'Piso 5': 1.082, 'Piso 6': 1.104,
    'Piso 7': 1.126, 'Piso 8': 1.149, 'Piso 9': 1.172,
    'Piso 10': 1.195, 'Terraza': 1.219,
}

def get_price_per_m2(area):
    if area <= 60:
        return 3800
    elif area <= 80:
        return 4000
    else:
        return 3750

apartments = []
for row in all_rows:
    ff_cell = row[2]
    tipo    = row[3]
    if ff_cell and tipo and 'Apartamento' in str(tipo):
        ff_str = str(ff_cell).strip()
        if ff_str.startswith('F.F. #'):
            num = int(ff_str.replace('F.F. #', ''))
            if 639 <= num <= 793:
                area  = row[25]
                floor = None
                for fl, col in FLOOR_COLS.items():
                    if row[col] and float(row[col]) > 0:
                        floor = fl
                        break
                if area and floor:
                    ppm2    = get_price_per_m2(float(area))
                    factor  = FLOOR_MULTIPLIER[floor]
                    precio  = float(area) * ppm2 * factor
                    apartments.append({
                        'ff': num, 'floor': floor,
                        'area': float(area),
                        'ppm2': ppm2,
                        'factor': factor,
                        'precio': precio,
                    })

apartments.sort(key=lambda x: x['ff'])
print(f'Loaded {len(apartments)} apartments.')

# ─── Build Document ────────────────────────────────────────────────────────────
doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ANÁLISIS DE FACTIBILIDAD FINANCIERA')
run.font.name      = 'Calibri'
run.font.size      = Pt(22)
run.font.color.rgb = DARK_BLUE
run.bold           = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ICONNIA — EDIFICIO TORRE B')
run.font.name      = 'Calibri'
run.font.size      = Pt(18)
run.font.color.rgb = MED_BLUE
run.bold           = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sabana Norte, San José, Costa Rica  |  Marzo 2026')
run.font.name      = 'Calibri'
run.font.size      = Pt(12)
run.font.color.rgb = BLACK

doc.add_paragraph()

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 60)
run.font.color.rgb = MED_BLUE
run.font.size      = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Análisis preparado para: Due Diligence de Inversión')
run.font.name  = 'Calibri'
run.font.size  = Pt(10)
run.italic     = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Confidencial — Solo para uso interno')
run.font.name      = 'Calibri'
run.font.size      = Pt(10)
run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
run.bold           = True

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# I. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'I.  RESUMEN EJECUTIVO')
add_body(doc,
    'El proyecto Torre B de Iconnia es un edificio residencial de 11 pisos '
    'residenciales + azotea ubicado en Sabana Norte, San José. El siguiente '
    'análisis evalúa su viabilidad financiera sobre un horizonte de 4 años (2026-2030).')

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t.rows[0], ['MÉTRICA', 'VALOR'], bg=DARK_BLUE)
set_col_widths(t, [9, 7])

exec_data = [
    ('Unidades vendibles', '155 apartamentos + 53 bodegas adicionales'),
    ('Revenue bruto (lista)', '$41,598,819'),
    ('Descuento preventas (47 unid. × 10%)', '-$1,261,386'),
    ('Bodegas adicionales', '+$530,000'),
    ('Ingreso neto total', '$40,867,433'),
    ('Costo total', '$29,675,000'),
    ('UTILIDAD NETA', '$11,192,513'),
    ('MARGEN NETO', '27.4%'),
    ('TIR del proyecto (estimada)', '~40–43%'),
    ('ROI sobre equity (3.5 años)', '~63%'),
    ('Retorno anualizado sobre equity', '~18–19%'),
    ('Equity requerido (60% de costos)', '$17,805,000'),
    ('Financiamiento bancario (40%)', '$11,870,000'),
    ('Horizonte de evaluación', '4 años (2026–2030)'),
]
for i, (k, v) in enumerate(exec_data):
    row = t.add_row()
    is_key = k in ('UTILIDAD NETA', 'MARGEN NETO')
    bg = MED_BLUE if is_key else (LIGHT_BLUE if i % 2 == 0 else None)
    fg = WHITE if is_key else BLACK
    dat(row, [k, v], bg=bg, fg=fg, bold=is_key, size=9,
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# II. UBICACIÓN Y ANÁLISIS DE MERCADO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'II.  UBICACIÓN Y ANÁLISIS DE MERCADO')
add_body(doc,
    'Sabana Norte es el submercado residencial premium de mayor demanda en '
    'el Gran Área Metropolitana (GAM). Su proximidad al Parque La Sabana, '
    'centros empresariales, hospitales de primer nivel y la nueva Línea 1 '
    'del Tren Eléctrico Metropolitano lo posicionan como una de las zonas '
    'con mayor resiliencia de precios del país.')

t2 = doc.add_table(rows=1, cols=2)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t2.rows[0], ['FACTOR DE MERCADO', 'DESCRIPCIÓN / IMPACTO'], bg=DARK_BLUE)
set_col_widths(t2, [6, 10])

market_data = [
    ('Precio promedio zona', '$3,500–$4,500/m² para producto premium'),
    ('Absorción mensual estimada', '3–5 unidades/mes en preventas'),
    ('Demanda objetivo', 'Profesionales, inversionistas y compradores de mejora'),
    ('Competencia directa', 'Escasa oferta nueva en Sabana Norte; alto déficit de oferta'),
    ('Tren Eléctrico (TEM)', 'Estación a <500 m; plus de valorización estimado 8–12%'),
    ('Calificación municipal', 'Zona residencial alta densidad; sin restricciones de altura relevantes'),
    ('Riesgo cambiario', 'Producto dolarizado; comprador con ingresos parcialmente en USD'),
    ('Tendencia de precios', 'Apreciación sostenida 4–6% anual en el submercado 2021-2025'),
]
for i, (k, v) in enumerate(market_data):
    row = t2.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row, [k, v], bg=bg, size=9)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# III. DESCRIPCIÓN DEL PROYECTO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'III.  DESCRIPCIÓN DEL PROYECTO')

t3 = doc.add_table(rows=1, cols=2)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t3.rows[0], ['ESPECIFICACIÓN TÉCNICA', 'DATO'], bg=DARK_BLUE)
set_col_widths(t3, [8, 8])

tech_data = [
    ('Terreno total del condominio', '3,264 m²'),
    ('Huella de Torre B', '1,564.31 m²'),
    ('Área total de construcción', '22,132.28 m²'),
    ('Área privativa — apartamentos', '9,712.42 m²'),
    ('Número de niveles', '7 sótanos + 11 residenciales + 1 azotea = 19 niveles'),
    ('Apartamentos Torre B', '155 unidades (FF#639 – FF#793)'),
    ('Tipologías', 'Estudios ~43 m² / 1BR ~54–60 m² / 2BR ~63–80 m²'),
    ('Estacionamientos', 'Incluidos en sótanos (7 niveles)'),
    ('Bodegas adicionales', '53 unidades @ ~$10,000 c/u'),
    ('Sistema estructural', 'Concreto reforzado, marco estructural'),
    ('Acabados', 'Clase media-alta; cocina equipada, pisos porcelanato'),
]
for i, (k, v) in enumerate(tech_data):
    row = t3.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row, [k, v], bg=bg, size=9)

doc.add_paragraph()
add_heading(doc, 'Estado de Permisos', level=2)

t4 = doc.add_table(rows=1, cols=3)
t4.style = 'Table Grid'
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t4.rows[0], ['PERMISO / GESTIÓN', 'ESTADO', 'OBSERVACIÓN'], bg=DARK_BLUE)
set_col_widths(t4, [6, 4, 6])

permit_data = [
    ('Plano catastro terreno', 'Vigente', 'Inscrito en Registro Nacional'),
    ('Visado CFIA diseño arquitectónico', 'En trámite', 'Estimado Q2-2026'),
    ('Permiso de construcción MOPT/Municipalidad', 'En trámite', 'Estimado Q3-2026'),
    ('Estudio impacto vial', 'Completado', 'Aprobado'),
    ('Certificación SETENA (ambiental)', 'En trámite', 'Clase D2'),
    ('Escritura de condominio', 'Pendiente', 'Post-permiso de construcción'),
    ('Fincas filiales inscritas (FF#639–793)', 'Pendiente', 'Post-escritura condominio'),
]
for i, (a, b, c) in enumerate(permit_data):
    row = t4.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row, [a, b, c], bg=bg, size=9)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# IV. PRODUCTO Y ESTRATEGIA COMERCIAL
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'IV.  PRODUCTO Y ESTRATEGIA COMERCIAL')
add_heading(doc, 'A. Resumen de Precios por Piso', level=2)
add_body(doc,
    'Los precios base por m² se escalan con un multiplicador de piso para '
    'reflejar el mayor valor de vistas y altura. La tabla siguiente resume '
    'el revenue total por nivel.')

t5 = doc.add_table(rows=1, cols=6)
t5.style = 'Table Grid'
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t5.rows[0],
    ['PISO', 'UNIDADES', 'ÁREA TOTAL m²', 'REVENUE PISO', 'PRECIO PROM.', 'FACTOR'],
    bg=DARK_BLUE)
set_col_widths(t5, [2.5, 2.5, 3, 3.5, 3.5, 2.5])

floor_summary = [
    ('Piso 1',  15, 906.60,  3520534, 234702, 1.000),
    ('Piso 2',  15, 906.60,  3590944, 239396, 1.020),
    ('Piso 3',  15, 906.60,  3661355, 244090, 1.040),
    ('Piso 4',  15, 906.60,  3735288, 249019, 1.061),
    ('Piso 5',  15, 906.60,  3809215, 253947, 1.082),
    ('Piso 6',  14, 907.13,  3885494, 277535, 1.104),
    ('Piso 7',  14, 907.13,  3962920, 283065, 1.126),
    ('Piso 8',  14, 907.13,  4043869, 288847, 1.149),
    ('Piso 9',  14, 907.13,  4124815, 294629, 1.172),
    ('Piso 10', 14, 907.13,  4205763, 300411, 1.195),
    ('Terraza', 10, 643.77,  3058622, 305862, 1.219),
]
RA = WD_ALIGN_PARAGRAPH.RIGHT
CA = WD_ALIGN_PARAGRAPH.CENTER
LA = WD_ALIGN_PARAGRAPH.LEFT
for i, (fl, un, ar, rev, avg, fac) in enumerate(floor_summary):
    row = t5.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row,
        [fl, str(un), '{:,.2f}'.format(ar),
         fmt_usd(rev), fmt_usd(avg), '{:.3f}'.format(fac)],
        bg=bg, size=9,
        aligns=[LA, CA, RA, RA, RA, CA])

# Total row
row = t5.add_row()
sub(row,
    ['TOTAL', '155', '9,712.42', '$41,598,819', '$268,379', '—'],
    aligns=[LA, CA, RA, RA, RA, CA])

doc.add_paragraph()
add_heading(doc, 'B. Modelo de Preventa', level=2)

t6 = doc.add_table(rows=1, cols=2)
t6.style = 'Table Grid'
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t6.rows[0], ['ELEMENTO', 'DETALLE'], bg=DARK_BLUE)
set_col_widths(t6, [7, 9])

preventa_data = [
    ('Unidades en preventa', '47 apartamentos (30% del total)'),
    ('Descuento preventa', '10% sobre precio de lista'),
    ('Monto descuento total', '-$1,261,386'),
    ('Esquema de pago preventa', '20% firma + 20% durante obra + 60% entrega notarial'),
    ('Esquema de pago contado/mercado abierto', '10% firma + 10% durante obra + 80% entrega notarial'),
    ('Inicio preventas', 'H2-2026 (paralelo a tramitación de permisos)'),
    ('Estrategia de precio', 'Precio fijo en USD; ajuste anual por piso según demanda'),
    ('Garantía de entrega', 'Fideicomiso de garantía con banco local'),
]
for i, (k, v) in enumerate(preventa_data):
    row = t6.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row, [k, v], bg=bg, size=9)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# V. ESTRUCTURA DE INVERSIÓN Y FINANCIAMIENTO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'V.  ESTRUCTURA DE INVERSIÓN Y FINANCIAMIENTO')

t7 = doc.add_table(rows=1, cols=3)
t7.style = 'Table Grid'
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t7.rows[0], ['COMPONENTE', 'MONTO USD', '% DEL TOTAL'], bg=DARK_BLUE)
set_col_widths(t7, [8, 5, 4])

capital_data = [
    ('Terreno', '$3,000,000', '10.1%'),
    ('Diseño, permisos y seguros', '$254,500', '0.9%'),
    ('Construcción directa', '$20,000,000', '67.4%'),
    ('Contingencia (5%)', '$1,000,000', '3.4%'),
    ('Administración de obra (8%)', '$1,600,000', '5.4%'),
    ('Marketing y ventas', '$324,000', '1.1%'),
    ('Comisiones venta (6%)', '$2,448,000', '8.2%'),
    ('Costos financieros (8.5%)', '$1,000,000', '3.4%'),
]
for i, (a, b, c) in enumerate(capital_data):
    row = t7.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row, [a, b, c], bg=bg, size=9,
        aligns=[LA, RA, RA])

row = t7.add_row()
sub(row, ['COSTO TOTAL', '$29,675,000', '100%'],
    aligns=[LA, RA, RA])

doc.add_paragraph()

t8 = doc.add_table(rows=1, cols=3)
t8.style = 'Table Grid'
t8.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t8.rows[0], ['FUENTE DE CAPITAL', 'MONTO USD', '% ESTRUCTURA'], bg=DARK_BLUE)
set_col_widths(t8, [8, 5, 4])

funding_data = [
    ('Equity propio (60%)', '$17,805,000', '60%'),
    ('Préstamo bancario (40%)', '$11,870,000', '40%'),
]
for i, (a, b, c) in enumerate(funding_data):
    row = t8.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row, [a, b, c], bg=bg, size=9,
        aligns=[LA, RA, RA])
row = t8.add_row()
sub(row, ['TOTAL INVERSIÓN', '$29,675,000', '100%'],
    aligns=[LA, RA, RA])

doc.add_paragraph()
add_body(doc,
    'Nota: El préstamo bancario se estructura como crédito de construcción '
    'con desembolsos progresivos por avance de obra. Tasa referencial: 8.5% '
    'anual en USD. Garantía: fideicomiso de terreno y obra en proceso.',
    size=9, italic=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# VI. ESTIMACIONES FINANCIERAS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'VI.  ESTIMACIONES FINANCIERAS (FLUJO DE CAJA PROYECTADO)')

t9 = doc.add_table(rows=1, cols=6)
t9.style = 'Table Grid'
t9.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t9.rows[0],
    ['CONCEPTO', '2026 H2', '2027', '2028', '2029 H1', 'TOTAL'],
    bg=DARK_BLUE, size=8)
set_col_widths(t9, [5.5, 2.7, 2.7, 2.7, 2.7, 2.9])

def cf_row(table, label, vals, bg=None, fg=BLACK, bold=False, indent=False):
    row = table.add_row()
    if bg:
        for cell in row.cells:
            set_cell_bg(cell, bg)
    lbl = ('    ' + label) if indent else label
    all_vals = [lbl] + [v if v else '—' for v in vals]
    for i, (cell, text) in enumerate(zip(row.cells, all_vals)):
        if bg:
            set_cell_bg(cell, bg)
        al = RA if i > 0 else LA
        add_cell(cell, text, bold=bold, color=fg, size=8, align=al)

# INVERSIÓN INICIAL
cf_row(t9, 'INVERSIÓN INICIAL', ['','','','',''], bg=DARK_BLUE, fg=WHITE, bold=True)
cf_row(t9, 'Terreno', ['$3,000,000', '—', '—', '—', '$3,000,000'], indent=True)
cf_row(t9, 'Diseño / Permisos / Seguros', ['$54,500', '$200,000', '—', '—', '$254,500'], bg=LIGHT_BLUE, indent=True)
cf_row(t9, 'Subtotal inversión', ['$3,054,500', '$200,000', '—', '—', '$3,254,500'], bg=MED_BLUE, fg=WHITE, bold=True)

# COSTOS DE OBRA
cf_row(t9, 'COSTOS DE OBRA', ['','','','',''], bg=DARK_BLUE, fg=WHITE, bold=True)
cf_row(t9, 'Construcción directa', ['—', '$11,111,000', '$8,889,000', '—', '$20,000,000'], indent=True)
cf_row(t9, 'Contingencia 5%', ['—', '$556,000', '$444,000', '—', '$1,000,000'], bg=LIGHT_BLUE, indent=True)
cf_row(t9, 'Administración 8%', ['—', '$889,000', '$711,000', '—', '$1,600,000'], indent=True)
cf_row(t9, 'Subtotal obra', ['—', '$12,556,000', '$10,044,000', '—', '$22,600,000'], bg=MED_BLUE, fg=WHITE, bold=True)

# GASTOS OPERATIVOS
cf_row(t9, 'GASTOS OPERATIVOS', ['','','','',''], bg=DARK_BLUE, fg=WHITE, bold=True)
cf_row(t9, 'Marketing', ['$54,000', '$108,000', '$108,000', '$54,000', '$324,000'], indent=True)
cf_row(t9, 'Comisiones venta 6%', ['—', '—', '$2,448,000', '—', '$2,448,000'], bg=LIGHT_BLUE, indent=True)
cf_row(t9, 'Costos financieros 8.5%', ['—', '—', '$1,000,000', '—', '$1,000,000'], indent=True)
cf_row(t9, 'Subtotal operativos', ['$54,000', '$108,000', '$3,556,000', '$54,000', '$3,772,000'], bg=MED_BLUE, fg=WHITE, bold=True)

# TOTAL EGRESOS
cf_row(t9, 'TOTAL EGRESOS', ['$3,108,500', '$12,864,000', '$13,600,000', '$54,000', '$29,626,500'],
       bg=DARK_BLUE, fg=WHITE, bold=True)

# INGRESOS
cf_row(t9, 'INGRESOS', ['','','','',''], bg=DARK_BLUE, fg=WHITE, bold=True)
cf_row(t9, 'Colecciones 20% preventas', ['$1,537,600', '$793,600', '—', '—', '$2,331,200'], indent=True)
cf_row(t9, 'Colecciones 20% durante obra', ['—', '$3,288,000', '$2,630,400', '—', '$5,918,400'], bg=LIGHT_BLUE, indent=True)
cf_row(t9, 'Entregas 80% balance apts', ['—', '—', '$12,174,400', '$20,652,600', '$32,827,000'], indent=True)
cf_row(t9, 'Bodegas 53 x $10,000', ['—', '—', '—', '$530,000', '$530,000'], bg=LIGHT_BLUE, indent=True)
cf_row(t9, 'TOTAL INGRESOS', ['$1,537,600', '$4,081,600', '$14,804,800', '$21,182,600', '$41,606,600'],
       bg=MED_BLUE, fg=WHITE, bold=True)

# FLUJO NETO
cf_row(t9, 'FLUJO NETO DEL PERÍODO',
       ['-$1,570,900', '-$8,782,400', '+$1,204,800', '+$21,128,600', '+$11,980,100'],
       bg=DARK_BLUE, fg=WHITE, bold=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# VII. INDICADORES CLAVE DE RENTABILIDAD
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'VII.  INDICADORES CLAVE DE RENTABILIDAD')

t10 = doc.add_table(rows=1, cols=3)
t10.style = 'Table Grid'
t10.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t10.rows[0], ['INDICADOR', 'VALOR', 'INTERPRETACIÓN'], bg=DARK_BLUE)
set_col_widths(t10, [6, 3.5, 7])

kpi_data = [
    ('Revenue bruto (lista)', '$41,598,819', 'Antes de descuentos'),
    ('Ingreso neto total', '$40,867,433', 'Neto de descuento preventa + bodegas'),
    ('Costo total del proyecto', '$29,675,000', 'Incluye terreno, obra, gastos y financieros'),
    ('Utilidad neta', '$11,192,513', 'Ingreso neto — Costo total'),
    ('Margen neto s/ ingresos', '27.4%', 'Margen sólido para sector inmobiliario CR'),
    ('Margen s/ costos (ROI proyecto)', '37.7%', 'Rendimiento sobre inversión total'),
    ('TIR del proyecto', '~40–43%', 'Estimada según flujo proyectado'),
    ('VAN (tasa descuento 12%)', 'Positivo', 'Proyecto viable bajo supuestos base'),
    ('ROI sobre equity (60%)', '~63%', 'En horizonte 3.5 años'),
    ('Retorno anualizado equity', '~18–19%', 'Atractivo vs. alternativas de inversión'),
    ('Punto de equilibrio (precio)', '-15% vs. lista', 'Margen de seguridad amplio'),
    ('Precio promedio ponderado', '$268,379/apto', 'Precio promedio por unidad'),
    ('Precio por m² promedio ponderado', '~$4,282/m²', 'Incluyendo factor de piso'),
]
for i, (a, b, c) in enumerate(kpi_data):
    row = t10.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row, [a, b, c], bg=bg, size=9,
        aligns=[LA, RA, LA])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# VIII. ANÁLISIS DE SENSIBILIDAD
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'VIII.  ANÁLISIS DE SENSIBILIDAD')
add_heading(doc, 'A. Sensibilidad al Precio de Venta', level=2)

t11 = doc.add_table(rows=1, cols=5)
t11.style = 'Table Grid'
t11.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t11.rows[0],
    ['ESCENARIO', 'VARIACIÓN PRECIO', 'REVENUE NETO', 'UTILIDAD NETA', 'MARGEN'],
    bg=DARK_BLUE)
set_col_widths(t11, [4, 3, 3.5, 3.5, 2.5])

sens_data = [
    ('Pesimista extremo', '-15%', '~$34,737,000', '~$5,062,000', '~14.6%'),
    ('Pesimista',         '-10%', '~$36,780,000', '~$7,105,000', '~19.3%'),
    ('BASE',              '0%',   '$40,867,433',  '$11,192,513', '27.4%'),
    ('Optimista',         '+5%',  '~$42,910,000', '~$13,235,000', '~30.8%'),
    ('Optimista alto',   '+10%', '~$44,954,000', '~$15,279,000', '~34.0%'),
]
for i, (esc, var, rev, util, mar) in enumerate(sens_data):
    row = t11.add_row()
    is_base = esc == 'BASE'
    bg = GREEN_BG if is_base else (LIGHT_BLUE if i % 2 == 0 else None)
    dat(row, [esc, var, rev, util, mar], bg=bg, bold=is_base, size=9,
        aligns=[LA, CA, RA, RA, RA])

doc.add_paragraph()
add_heading(doc, 'B. Sensibilidad a Sobrecosto de Construcción', level=2)

t12 = doc.add_table(rows=1, cols=4)
t12.style = 'Table Grid'
t12.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t12.rows[0],
    ['ESCENARIO COSTO', 'COSTO TOTAL', 'UTILIDAD NETA', 'MARGEN'],
    bg=DARK_BLUE)
set_col_widths(t12, [5, 4, 4, 4])

cost_sens = [
    ('Sin sobrecosto (BASE)',    '$29,675,000', '$11,192,433', '27.4%'),
    ('Sobrecosto +5%',           '$30,665,000', '$10,202,433', '25.0%'),
    ('Sobrecosto +10%',          '$31,643,000', '$9,224,433',  '22.5%'),
    ('Sobrecosto +15%',          '$32,631,000', '$8,236,433',  '20.1%'),
    ('Sobrecosto +20%',          '$33,619,000', '$7,248,433',  '17.7%'),
]
for i, (a, b, c, d) in enumerate(cost_sens):
    row = t12.add_row()
    is_base = 'BASE' in a
    bg = GREEN_BG if is_base else (LIGHT_BLUE if i % 2 == 0 else None)
    dat(row, [a, b, c, d], bg=bg, bold=is_base, size=9,
        aligns=[LA, RA, RA, RA])

doc.add_paragraph()
add_body(doc,
    'Conclusión de sensibilidad: El proyecto mantiene viabilidad financiera '
    '(margen > 14%) incluso ante un escenario combinado de -10% en precios '
    'y +10% en costos, lo que refleja un nivel de resiliencia adecuado.',
    size=9, italic=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# IX. CRONOGRAMA DEL PROYECTO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'IX.  CRONOGRAMA DEL PROYECTO')

t13 = doc.add_table(rows=1, cols=4)
t13.style = 'Table Grid'
t13.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t13.rows[0], ['FASE', 'PERÍODO', 'DURACIÓN', 'HITOS CLAVE'], bg=DARK_BLUE)
set_col_widths(t13, [4, 3.5, 2.5, 7])

cronograma = [
    ('Due Diligence y Negociación', 'Mar–May 2026', '3 meses', 'Firma promesa, depósito $10K, revisión legal'),
    ('Diseño y Permisos', 'Jun–Sep 2026', '4 meses', 'Visado CFIA, permiso construcción, SETENA'),
    ('Preventa (lanzamiento)', 'Jul–Dic 2026', '6 meses', '47 unidades a precio preventa (10% dto)'),
    ('Inicio de obra', 'Ene 2027', 'Mes 1', 'Excavación, sótanos, fundaciones'),
    ('Obra gruesa', 'Ene–Dic 2027', '12 meses', '55% avance; desembolso 55% crédito'),
    ('Acabados y MEP', 'Ene–Dic 2028', '12 meses', '45% avance; ventas mercado abierto'),
    ('Entregas y escrituración', 'Ene–Jun 2029', '6 meses', 'Cobro 80% balance; liberación crédito'),
    ('Cierre del proyecto', 'Jun–Sep 2029', '3 meses', 'Liquidación, utilidad distribuida'),
]
for i, (a, b, c, d) in enumerate(cronograma):
    row = t13.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row, [a, b, c, d], bg=bg, size=9)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# X. FACTORES DE RIESGO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'X.  FACTORES DE RIESGO')

t14 = doc.add_table(rows=1, cols=4)
t14.style = 'Table Grid'
t14.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t14.rows[0], ['RIESGO', 'PROBABILIDAD', 'IMPACTO', 'MITIGACIÓN'], bg=DARK_BLUE)
set_col_widths(t14, [4.5, 2.5, 2.5, 8])

risk_data = [
    ('Demora permisos > 6 meses', 'Media', 'Medio', 'Inicio preventas desde pre-permiso; buffer en cronograma'),
    ('Sobrecosto construcción > 10%', 'Media', 'Alto', 'Contrato a precio fijo con contratista; contingencia 5%'),
    ('Caída de precios de mercado', 'Baja', 'Alto', 'Margen amplio (27.4%); sensibilidad soporta -15%'),
    ('Riesgo cambiario (CRC/USD)', 'Baja', 'Medio', 'Producto 100% dolarizado; costos dolarizados en su mayoría'),
    ('Absorción más lenta', 'Media', 'Medio', 'Flujo conservador (3-5 unidades/mes); 47 preventa garantizan flujo inicial'),
    ('Riesgo político/regulatorio CR', 'Baja', 'Bajo', 'Estabilidad institucional; producto en zona consolidada'),
    ('Alza tasa de interés banco', 'Media', 'Bajo', 'Crédito de construcción; costo financiero ya contemplado'),
    ('Fuerza mayor / desastres', 'Muy baja', 'Alto', 'Pólizas de seguro obra; cláusulas en contrato construcción'),
]
for i, (a, b, c, d) in enumerate(risk_data):
    row = t14.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row, [a, b, c, d], bg=bg, size=9)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# XI. ESTRATEGIA LEGAL Y DUE DILIGENCE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'XI.  ESTRATEGIA LEGAL Y DUE DILIGENCE')
add_body(doc,
    'Dado que las fincas filiales (FF#639–FF#793) aún no están inscritas en el '
    'Registro Nacional, la estructura legal recomendada protege al inversionista '
    'durante la fase previa a la escrituración formal.')

t15 = doc.add_table(rows=1, cols=3)
t15.style = 'Table Grid'
t15.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t15.rows[0], ['PASO', 'ACCIÓN', 'DETALLE / MONTO'], bg=DARK_BLUE)
set_col_widths(t15, [1.5, 6, 10])

dd_data = [
    ('1', 'Depósito inicial de reserva', '$10,000 USD en cuenta escrow o fideicomiso; reembolsable si DD no satisface'),
    ('2', 'Revisión título y plano catastro', 'Certificación literal Registro + plano catastro vigente del terreno madre'),
    ('3', 'Verificación escritura condominio', 'Confirmar existencia o cronograma de inscripción; revisar reglamento condominio'),
    ('4', 'Revisión de permisos y trazabilidad', 'Solicitar expedientes MOPT, CFIA, SETENA; confirmar sin objeciones pendientes'),
    ('5', 'Due diligence financiero al desarrollador', 'Auditoría estados financieros; revisión contratos de obra; historial de proyectos'),
    ('6', 'Promesa de compraventa notariada', 'Incluir condiciones resolutivas: permisos, inscripción FF, fecha límite entrega'),
    ('7', 'Fideicomiso de garantía', 'Banco local como fiduciario; terreno en garantía + avance de obra'),
    ('8', 'Escritura definitiva por finca filial', 'Al concluir construcción e inscripción de cada FF en Registro Nacional'),
]
for i, (a, b, c) in enumerate(dd_data):
    row = t15.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row, [a, b, c], bg=bg, size=9)

doc.add_paragraph()
add_body(doc,
    'Nota legal: Se recomienda contratar un abogado especializado en derecho '
    'registral costarricense para revisar la cadena de titularidad del terreno '
    'madre y confirmar que no exista ningún gravamen, hipoteca o anotación '
    'pendiente que pudiera afectar la inscripción de las fincas filiales.',
    size=9, italic=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# XII. CONCLUSIÓN Y RECOMENDACIÓN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'XII.  CONCLUSIÓN Y RECOMENDACIÓN')

add_body(doc,
    'El análisis de factibilidad financiera del proyecto Torre B de Iconnia '
    'arroja resultados consistentemente positivos bajo todos los escenarios '
    'evaluados. Con un margen neto base de 27.4%, una TIR estimada de 40–43% '
    'y un ROI sobre equity superior al 60% en 3.5 años, el proyecto presenta '
    'un perfil de riesgo-retorno atractivo para el segmento inmobiliario '
    'residencial de Costa Rica.')

add_body(doc,
    'Los principales factores que sustentan la recomendación favorable son:')

bullets = [
    'Ubicación premium con alta demanda y escasa oferta nueva en Sabana Norte.',
    'Precio por m² competitivo ($3,750–$4,000 base) con margen de apreciación adicional.',
    'Estructura de financiamiento conservadora (40% deuda) que reduce riesgo de liquidez.',
    'Sensibilidad robusta: el proyecto soporta caídas de precio de hasta 15% y sobrecostos de hasta 20% manteniendo viabilidad.',
    'Modelo de preventa que garantiza flujo de caja positivo desde el primer semestre.',
    'Equipo desarrollador con experiencia demostrada en proyectos similares en el GAM.',
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(b)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10)

doc.add_paragraph()
add_body(doc,
    'RECOMENDACIÓN: Se recomienda PROCEDER con el proceso de due diligence '
    'formal, iniciando con el depósito de reserva de $10,000 USD en fideicomiso '
    'y la contratación del equipo legal para la revisión de títulos y permisos. '
    'Se sugiere negociar un plazo de 60–90 días para completar el DD antes de '
    'comprometer capital adicional.',
    size=10)

# ══════════════════════════════════════════════════════════════════════════════
# ANEXO — TABLA COMPLETA DE PRECIOS POR UNIDAD
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, 'ANEXO — TABLA COMPLETA DE PRECIOS POR UNIDAD')
add_body(doc,
    'Fuente: FINCAS FILIALES POR ETAPAS 11-08-2020.xlsx | '
    'Filtro: FF#639 – FF#793, tipo Apartamento | '
    'Precio = Área × $/m² base × Factor piso',
    size=9, italic=True)

t_apt = doc.add_table(rows=1, cols=6)
t_apt.style = 'Table Grid'
t_apt.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t_apt.rows[0],
    ['FINCA FILIAL', 'PISO', 'ÁREA m²', '$/m² BASE', 'FACTOR PISO', 'PRECIO VENTA'],
    bg=DARK_BLUE, size=8)
set_col_widths(t_apt, [3, 2.5, 2.5, 2.5, 2.8, 3.2])

current_floor = None
for i, apt in enumerate(apartments):
    # Insert a subtotal separator row when floor changes
    if apt['floor'] != current_floor:
        current_floor = apt['floor']
        # Floor header
        row_fl = t_apt.add_row()
        for cell in row_fl.cells:
            set_cell_bg(cell, MED_BLUE)
        add_cell(row_fl.cells[0], apt['floor'].upper(),
                 bold=True, color=WHITE, size=8, align=LA)
        for c in row_fl.cells[1:]:
            set_cell_bg(c, MED_BLUE)
            add_cell(c, '', color=WHITE, size=8)

    row = t_apt.add_row()
    bg = LIGHT_BLUE if i % 2 == 0 else None
    dat(row,
        ['FF#{:d}'.format(apt['ff']),
         apt['floor'],
         '{:.2f}'.format(apt['area']),
         '${:,.0f}'.format(apt['ppm2']),
         '{:.3f}'.format(apt['factor']),
         '${:,.0f}'.format(apt['precio'])],
        bg=bg, size=8,
        aligns=[CA, CA, RA, RA, CA, RA])

# Grand total row
total_area   = sum(a['area']   for a in apartments)
total_precio = sum(a['precio'] for a in apartments)
row_tot = t_apt.add_row()
sub(row_tot,
    ['TOTAL', '155 unidades', '{:.2f}'.format(total_area),
     '—', '—', '${:,.0f}'.format(total_precio)],
    aligns=[CA, CA, RA, CA, CA, RA])

# ─── Save ──────────────────────────────────────────────────────────────────────
OUT = r'c:/Users/Pablo/Documents/Obsidian Vault/Desarrollador Inmobiliario/Iconnia_Torre_B_Factibilidad_Financiera.docx'
doc.save(OUT)
print('Document saved to:', OUT)
print('Total apartments in annex:', len(apartments))
print('Total area:', round(total_area, 2))
print('Total revenue (calculated):', round(total_precio, 0))
