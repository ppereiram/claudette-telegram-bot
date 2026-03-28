"""
Tools Registry para Claudette Bot.
Define TODOS los schemas de herramientas y su ejecución.
Portado completo desde bot.py monolítico.
"""

import os
import tempfile
import logging
import requests
from datetime import datetime
from config import OPENAI_API_KEY, OPENWEATHER_API_KEY, DEFAULT_LOCATION, FIRECRAWL_API_KEY, logger
from memory_manager import save_fact, get_fact
from library import search_library, search_by_author, search_by_tag, get_book_content, get_library_stats
from knowledge_base import KB_TOOLS_SCHEMA, execute_kb_tool

# --- Imports de servicios Google ---
import google_calendar
import gmail_service
import google_tasks
import google_drive
import google_contacts
import google_places

# --- Clients opcionales ---
from openai import OpenAI
openai_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

# --- Imports para libros ---
try:
    import pypdf
except ImportError:
    pypdf = None

# --- Import para generación de documentos ---
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    DocxDocument = None

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    openpyxl = None

try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
except ImportError:
    ebooklib = None

# --- Búsqueda web ---
try:
    from googlesearch import search as google_search_func
except ImportError:
    google_search_func = None


# =====================================================
# FUNCIONES DE APOYO
# =====================================================

def clean_date_iso(date_str, is_end=False):
    """Asegura formato ISO con timezone Costa Rica (UTC-6)."""
    if 'T' not in date_str:
        time_part = "T23:59:59" if is_end else "T00:00:00"
        return f"{date_str}{time_part}-06:00"
    # Si tiene T pero no tiene timezone, agregar -06:00
    if '+' not in date_str and '-06:00' not in date_str and not date_str.endswith('Z'):
        return f"{date_str}-06:00"
    # Si viene con Z (UTC), reemplazar por -06:00 — Claude a veces manda UTC
    if date_str.endswith('Z'):
        return date_str[:-1] + "-06:00"
    return date_str


def get_weather(lat, lon):
    """Obtener clima de OpenWeather."""
    if not OPENWEATHER_API_KEY:
        return "⚠️ No tengo configurada la API Key de OpenWeather."
    try:
        url = f"https://api.openweathermap.org/data/2.5/weather?lat={lat}&lon={lon}&appid={OPENWEATHER_API_KEY}&units=metric&lang=es"
        res = requests.get(url, timeout=10).json()
        if res.get('cod') != 200:
            return f"Error clima: {res.get('message')}"
        desc = res['weather'][0]['description']
        temp = res['main']['temp']
        hum = res['main']['humidity']
        city = res['name']
        return f"ðŸŒ¦️ El clima en {city}: {desc.capitalize()}, {temp}°C, Humedad {hum}%."
    except Exception as e:
        return f"Error obteniendo clima: {e}"


def get_weather_by_city(city_name: str) -> str:
    """Obtener clima de OpenWeather por nombre de ciudad o pais."""
    if not OPENWEATHER_API_KEY:
        return "No tengo la API Key de OpenWeather."
    try:
        url = f"https://api.openweathermap.org/data/2.5/weather?q={city_name}&appid={OPENWEATHER_API_KEY}&units=metric&lang=es"
        res = requests.get(url, timeout=10).json()
        if res.get('cod') != 200:
            return f"No encontre el clima para '{city_name}': {res.get('message')}"
        desc = res['weather'][0]['description']
        temp = res['main']['temp']
        temp_min = res['main'].get('temp_min', temp)
        temp_max = res['main'].get('temp_max', temp)
        hum = res['main']['humidity']
        city = res['name']
        country = res.get('sys', {}).get('country', '')
        wind = res.get('wind', {}).get('speed', 0)
        return ("Clima en " + city + ", " + country + ": " + desc.capitalize() + chr(10) + "Temperatura: " + str(round(temp,1)) + "C (min " + str(round(temp_min)) + " / max " + str(round(temp_max)) + ")" + chr(10) + "Humedad: " + str(hum) + "% | Viento: " + str(wind) + " m/s")
    except Exception as e:
        return f"Error obteniendo clima: {e}"


def search_web_google(query, max_results=5):
    """Busca en web con fallback DuckDuckGo."""
    if google_search_func:
        try:
            results = []
            for result in google_search_func(query, num_results=max_results, advanced=True, lang="es"):
                results.append(f"ðŸ“° {result.title}\nðŸ”— {result.url}\nðŸ“ {result.description}\n")
            if results:
                return "\n".join(results)
        except Exception as e:
            logger.warning(f"Google Search falló: {e}, usando DuckDuckGo...")

    try:
        from duckduckgo_search import DDGS
        ddgs = DDGS()
        results = []
        for r in ddgs.text(query, max_results=max_results, region="es-es"):
            results.append(f"ðŸ“° {r['title']}\nðŸ”— {r['href']}\nðŸ“ {r['body']}\n")
        if results:
            return "\n".join(results)
    except Exception as e:
        logger.error(f"DuckDuckGo también falló: {e}")

    return "No se encontraron resultados."


def search_news(topics=None):
    """
    Busca noticias usando RSS feeds directos + DuckDuckGo como fallback.
    Fuentes RSS: Reuters, BBC Mundo, El Pais, Hacker News, Financial Times
    """
    import urllib.request
    import xml.etree.ElementTree as ET
    from datetime import datetime, timezone

    # RSS feeds por categoria - mas confiables que DuckDuckGo
    RSS_FEEDS = {
        "INTELIGENCIA ARTIFICIAL": [
            "https://hnrss.org/frontpage?points=100",
            "https://feeds.feedburner.com/oreilly/radar",
        ],
        "GEOPOLITICA": [
            "https://feeds.bbci.co.uk/mundo/internacional/rss.xml",
            "https://rss.nytimes.com/services/xml/rss/nyt/World.xml",
        ],
        "MERCADOS/ECONOMIA": [
            "https://rss.nytimes.com/services/xml/rss/nyt/Business.xml",
            "https://feeds.bbci.co.uk/mundo/economia/rss.xml",
        ],
        "CIENCIA/TECNOLOGIA": [
            "https://rss.nytimes.com/services/xml/rss/nyt/Science.xml",
            "https://www.sciencedaily.com/rss/top/science.xml",
        ],
    }

    all_news = []
    headers = {"User-Agent": "Mozilla/5.0"}

    for category, feeds in RSS_FEEDS.items():
        cat_items = []
        for feed_url in feeds:
            if len(cat_items) >= 2:
                break
            try:
                req = urllib.request.Request(feed_url, headers=headers)
                with urllib.request.urlopen(req, timeout=5) as resp:
                    xml_data = resp.read()
                root = ET.fromstring(xml_data)
                ns = ""
                items = root.findall(".//item")
                for item in items[:2]:
                    title = item.findtext("title", "").strip()
                    link = item.findtext("link", "").strip()
                    if title and link and len(title) > 10:
                        cat_items.append("  - " + title + " | " + link)
                        if len(cat_items) >= 2:
                            break
            except Exception:
                continue

        if cat_items:
            all_news.append(f"{category}:\n" + "\n".join(cat_items))

    # Si RSS falla, intentar DuckDuckGo
    if not all_news:
        try:
            from config import NEWS_TOPICS
            from duckduckgo_search import DDGS
            ddgs = DDGS()
            for topic in (topics or NEWS_TOPICS):
                clean_topic = topic.replace(" noticias", "").strip().upper()
                try:
                    results = ddgs.news(topic, max_results=2, region="es-es")
                    lines = [f"  - {r['title']}" for r in results]
                    if lines:
                        all_news.append(f"{clean_topic}:\n" + "\n".join(lines))
                except Exception:
                    continue
        except Exception:
            pass

    if not all_news:
        return "[NOTICIAS NO DISPONIBLES — los RSS fallaron hoy. No menciones noticias en el boletin, omite esa sección.]"

    return "\n\n".join(all_news)
def extract_text_from_pdf(file_path):
    """Extrae texto de PDF."""
    if not pypdf:
        return "Error: pypdf no está instalado."
    text = ""
    try:
        reader = pypdf.PdfReader(file_path)
        max_pages = min(len(reader.pages), 50)
        for i in range(max_pages):
            text += reader.pages[i].extract_text() + "\n"
    except Exception as e:
        logger.error(f"PDF Error: {e}")
    return text


def extract_text_from_epub(file_path):
    """Extrae texto de EPUB."""
    if not ebooklib:
        return "Error: ebooklib no está instalado."
    text = ""
    try:
        book = epub.read_epub(file_path, options={'ignore_ncx': True})
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            chapter_text = soup.get_text(separator='\n', strip=True)
            if chapter_text.strip():
                text += chapter_text + "\n\n"
    except Exception as e:
        logger.error(f"EPUB Error: {e}")
        text = f"Error leyendo EPUB: {e}"
    return text


def read_book_from_drive(query):
    """Busca y lee documentos desde Drive — soporta Google Docs, PDF, EPUB, TXT, MD."""
    service = google_drive.get_drive_service()
    if not service:
        return "Error conectando a Drive."
    try:
        results = service.files().list(
            q=f"name contains '{query}' and mimeType != 'application/vnd.google-apps.folder'",
            pageSize=1,
            fields="files(id, name, mimeType)"
        ).execute()
        items = results.get('files', [])
        if not items:
            return "No hallé el archivo."

        file_id = items[0]['id']
        file_name = items[0]['name']
        mime_type = items[0].get('mimeType', '')

        content = ""

        # Google Docs → exportar como texto plano
        if mime_type == 'application/vnd.google-apps.document':
            export = service.files().export(fileId=file_id, mimeType='text/plain').execute()
            content = export.decode('utf-8') if isinstance(export, bytes) else str(export)

        else:
            # Archivos binarios (PDF, EPUB, TXT, MD)
            request = service.files().get_media(fileId=file_id)
            file_content = request.execute()

            with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{file_name}") as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name

            lower_name = file_name.lower()
            if lower_name.endswith('.pdf'):
                content = extract_text_from_pdf(temp_path)
            elif lower_name.endswith('.epub'):
                content = extract_text_from_epub(temp_path)
            elif lower_name.endswith(('.txt', '.md')):
                with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            else:
                content = f"Formato '{file_name.split('.')[-1]}' no soportado. Válidos: Google Docs, PDF, EPUB, TXT, MD."

            os.unlink(temp_path)

        if not content.strip():
            return f"📖 Encontré '{file_name}' pero no pude extraer texto."
        if len(content) > 8000:
            return f"📖 {file_name} (Primeras páginas):\n{content[:8000]}\n\n[... Truncado. Pídeme una sección específica.]"
        return f"📖 {file_name}:\n{content}"

    except Exception as e:
        return f"Error leyendo libro: {e}"


def read_local_file(filename):
    """Lee archivos locales del sistema de modelos mentales y prompts."""
    ALLOWED_FILES = [
        'MODELS_DEEP.md', 'FRAMEWORK.md', 'ANTIPATTERNS.md', 'TEMPLATES.md',
        'CLAUDETTE_CORE.md', 'user_profile.md'
    ]

    if filename not in ALLOWED_FILES:
        return f"Archivo '{filename}' no permitido. Archivos válidos: {', '.join(ALLOWED_FILES)}"

    # Buscar en varias rutas posibles
    possible_paths = [
        f'prompts/{filename}',
        filename,
        f'/opt/render/project/src/{filename}',
        f'/opt/render/project/src/prompts/{filename}',
    ]

    for path in possible_paths:
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read()
                if len(content) > 15000:
                    return f"📖 {filename} (truncado a 15000 chars):\n{content[:15000]}\n\n[... Truncado. Pide una sección específica.]"
                return f"📖 {filename}:\n{content}"
            except Exception as e:
                return f"Error leyendo {filename}: {e}"

    return f"Archivo '{filename}' no encontrado."


def _fetch_with_firecrawl(url):
    """Usa Firecrawl API para extraer contenido de páginas que bloquean scraping directo."""
    if not FIRECRAWL_API_KEY:
        return None
    try:
        resp = requests.post(
            "https://api.firecrawl.dev/v1/scrape",
            headers={
                "Authorization": f"Bearer {FIRECRAWL_API_KEY}",
                "Content-Type": "application/json"
            },
            json={"url": url, "formats": ["markdown"]},
            timeout=30
        )
        if resp.status_code == 200:
            data = resp.json()
            if data.get("success"):
                md = data.get("data", {}).get("markdown", "")
                title = data.get("data", {}).get("metadata", {}).get("title", "")
                if md:
                    content = md[:8000]
                    if title:
                        return f"📰 **{title}**\n\n{content}"
                    return content
    except Exception as e:
        logger.warning(f"Firecrawl error: {e}")
    return None


def fetch_url(url):
    """Lee el contenido de una página web, tweet, artículo, etc."""
    import re as re_mod
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'es,en;q=0.9',
        }

        clean_url = url.strip()

        # Manejar URLs de X/Twitter
        if 'x.com/' in clean_url or 'twitter.com/' in clean_url:
            fx_url = clean_url.replace('x.com/', 'api.fxtwitter.com/').replace('twitter.com/', 'api.fxtwitter.com/')
            try:
                resp = requests.get(fx_url, headers=headers, timeout=15)
                if resp.status_code == 200:
                    import json
                    data = resp.json()
                    tweet = data.get('tweet', {})
                    author = tweet.get('author', {}).get('name', 'Unknown')
                    handle = tweet.get('author', {}).get('screen_name', '')
                    text = tweet.get('text', '')
                    if text:
                        media = tweet.get('media', {})
                        photos = media.get('photos', []) if media else []
                        result = f"Tweet de {author} (@{handle}):\n{text}"
                        if photos:
                            result += f"\n[{len(photos)} imagen(es) adjunta(s)]"
                        return result
            except Exception:
                pass
            # X.com bloquea scraping — no caer al fetch generico
            return (
                "No pude leer el tweet (X.com bloquea bots).\n"
                "Pega el texto del tweet en el chat y lo analizo."
            )


        # Manejar URLs de YouTube - extraer transcript
        if 'youtube.com/watch' in clean_url or 'youtu.be/' in clean_url:
            try:
                from youtube_transcript_api import YouTubeTranscriptApi
                import re as re_yt
                # Extraer video ID
                vid_match = re_yt.search(r'(?:v=|youtu\.be/)([\w-]{11})', clean_url)
                if vid_match:
                    video_id = vid_match.group(1)
                    # Intentar transcript en espanol, luego ingles, luego cualquiera
                    transcript_list = None
                    for lang in [['es'], ['en'], None]:
                        try:
                            if lang:
                                transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=lang)
                            else:
                                transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
                            break
                        except Exception:
                            continue
                    if transcript_list:
                        full_text = " ".join([t["text"] for t in transcript_list])
                        if len(full_text) > 50000:
                            full_text = full_text[:50000] + "\n\n[... Transcript truncado — video muy largo]"
                        # Obtener titulo via oEmbed
                        try:
                            oembed = requests.get(
                                "https://www.youtube.com/oembed",
                                params={"url": clean_url, "format": "json"},
                                timeout=5
                            ).json()
                            yt_title = oembed.get("title", "Video de YouTube")
                            yt_author = oembed.get("author_name", "")
                        except Exception:
                            yt_title = "Video de YouTube"
                            yt_author = ""
                        result = f"\U0001f4fa **{yt_title}**"
                        if yt_author:
                            result += f"\nCanal: {yt_author}"
                        result += f"\n\n**Transcript:**\n{full_text}"
                        return result
                    else:
                        return "No hay transcript disponible para este video de YouTube."
            except ImportError:
                pass
            except Exception as e:
                pass
            # Fallback: extraer metadata basica
            try:
                oembed = requests.get(
                    "https://www.youtube.com/oembed",
                    params={"url": clean_url, "format": "json"},
                    timeout=5
                ).json()
                return f"\U0001f4fa **{oembed.get('title', '')}**\nCanal: {oembed.get('author_name', '')}\n\n(Transcript no disponible para este video)"
            except Exception:
                pass

        resp = requests.get(clean_url, headers=headers, timeout=15, allow_redirects=True)
        resp.raise_for_status()

        content_type = resp.headers.get('content-type', '')

        if 'json' in content_type:
            import json
            return json.dumps(resp.json(), indent=2, ensure_ascii=False)[:6000]

        html = resp.text

        # Extraer título
        title_match = re_mod.search(r'<title[^>]*>(.*?)</title>', html, re_mod.IGNORECASE | re_mod.DOTALL)
        title = title_match.group(1).strip() if title_match else ''

        # Extraer meta description / og:description
        description = ''
        og_match = re_mod.search(r'<meta[^>]*property=["\']og:description["\'][^>]*content=["\'](.*?)["\']', html, re_mod.IGNORECASE)
        if og_match:
            description = og_match.group(1).strip()
        else:
            desc_match = re_mod.search(r'<meta[^>]*name=["\']description["\'][^>]*content=["\'](.*?)["\']', html, re_mod.IGNORECASE)
            if desc_match:
                description = desc_match.group(1).strip()

        # Limpiar HTML → texto
        text = html
        text = re_mod.sub(r'<script[^>]*>.*?</script>', '', text, flags=re_mod.DOTALL | re_mod.IGNORECASE)
        text = re_mod.sub(r'<style[^>]*>.*?</style>', '', text, flags=re_mod.DOTALL | re_mod.IGNORECASE)
        text = re_mod.sub(r'<nav[^>]*>.*?</nav>', '', text, flags=re_mod.DOTALL | re_mod.IGNORECASE)
        text = re_mod.sub(r'<footer[^>]*>.*?</footer>', '', text, flags=re_mod.DOTALL | re_mod.IGNORECASE)
        text = re_mod.sub(r'<br\s*/?>', '\n', text)
        text = re_mod.sub(r'</p>', '\n\n', text)
        text = re_mod.sub(r'</div>', '\n', text)
        text = re_mod.sub(r'</h[1-6]>', '\n\n', text)
        text = re_mod.sub(r'<[^>]+>', '', text)
        text = re_mod.sub(r'\n\s*\n', '\n\n', text)
        text = re_mod.sub(r' +', ' ', text)
        text = text.strip()

        if len(text) > 5000:
            text = text[:5000] + "\n\n[... Contenido truncado]"

        result = ""
        if title:
            result += f"ðŸ“° **{title}**\n"
        if description:
            result += f"_{description}_\n\n"
        result += text

        # Si el contenido extraido es pobre, intentar con Firecrawl
        if len(result.strip()) < 200:
            fc_result = _fetch_with_firecrawl(clean_url)
            if fc_result:
                return fc_result
        return result if result.strip() else "No pude extraer contenido legible de esa URL."

    except requests.exceptions.Timeout:
        fc_result = _fetch_with_firecrawl(url.strip())
        if fc_result:
            return fc_result
        return "⚠️ La pagina tardo demasiado en responder."
    except requests.exceptions.HTTPError as e:
        if e.response.status_code in (403, 401, 429):
            fc_result = _fetch_with_firecrawl(url.strip())
            if fc_result:
                return fc_result
        return f"⚠️ Error HTTP {e.response.status_code} al acceder a la URL."
    except Exception as e:
        fc_result = _fetch_with_firecrawl(url.strip())
        if fc_result:
            return fc_result
        return f"⚠️ Error leyendo URL: {e}"


# =====================================================
# GENERACIÃ“N DE DOCUMENTOS
# =====================================================

def generate_document(title, content, doc_format="docx"):
    """
    Genera un documento descargable (.docx o .md) a partir del contenido.
    Parsea formato Markdown básico para crear documentos Word con estilo.
    Retorna la ruta del archivo temporal generado.
    """
    import re as re_mod

    # Sanitizar título para nombre de archivo
    safe_title = re_mod.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:50]
    timestamp = datetime.now().strftime("%Y%m%d")

    if doc_format == "md":
        # Markdown: guardar directamente
        filename = f"{safe_title}_{timestamp}.md"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"# {title}\n\n")
            f.write(content)
        return filepath, filename

    # --- DOCX ---
    if not DocxDocument:
        # Fallback a markdown si python-docx no está
        logger.warning("python-docx no instalado, generando .md")
        return generate_document(title, content, "md")

    filename = f"{safe_title}_{timestamp}.docx"
    filepath = os.path.join(tempfile.gettempdir(), filename)

    doc = DocxDocument()

    # --- Estilos del documento ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 30, 30)

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Título principal
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtítulo con fecha y autor
    date_str = datetime.now().strftime("%d de %B, %Y")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Pablo Pereiram — {date_str}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.font.italic = True

    doc.add_paragraph("")  # Espacio

    # --- Parsear contenido Markdown → DOCX ---
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Línea vacía
        if not stripped:
            i += 1
            continue

        # Encabezados
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)

        # Separador horizontal
        elif stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.add_run('â”€' * 50).font.color.rgb = RGBColor(180, 180, 180)

        # Listas con viñetas
        elif stripped.startswith(('- ', 'â€¢ ', '* ')):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)

        # Listas numeradas
        elif re_mod.match(r'^\d+[\.\)]\s', stripped):
            text = re_mod.sub(r'^\d+[\.\)]\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)

        # Citas (blockquote)
        elif stripped.startswith('> '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        # Párrafos normales
        else:
            # Acumular líneas consecutivas como un solo párrafo
            para_lines = [stripped]
            while i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if (next_line and
                    not next_line.startswith(('#', '-', 'â€¢', '*', '>', '---', '***', '___')) and
                    not re_mod.match(r'^\d+[\.\)]\s', next_line)):
                    para_lines.append(next_line)
                    i += 1
                else:
                    break
            full_text = ' '.join(para_lines)
            p = doc.add_paragraph()
            _add_formatted_text(p, full_text)

        i += 1

    doc.save(filepath)
    return filepath, filename


def _add_formatted_text(paragraph, text):
    """Agrega texto con formato inline (bold, italic) a un párrafo docx."""
    import re as re_mod

    # Patrón unificado: ***bold+italic***, **bold**, *italic*
    pattern = re_mod.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*')

    last_end = 0
    for match in pattern.finditer(text):
        # Texto plano antes del match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        # Determinar tipo de formato
        if match.group(1):  # ***bold+italic***
            run = paragraph.add_run(match.group(1))
            run.bold = True
            run.italic = True
        elif match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True

        last_end = match.end()

    # Texto restante después del último match
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


# =====================================================
# GENERACIÃ“N DE SPREADSHEETS (EXCEL)
# =====================================================

def generate_spreadsheet(title, sheets_data):
    """
    Genera un archivo Excel (.xlsx) con múltiples hojas.
    
    sheets_data: lista de dicts con:
      - sheet_name: nombre de la hoja
      - headers: lista de strings con los encabezados
      - rows: lista de listas con los datos
    
    Retorna (filepath, filename).
    """
    import re as re_mod

    if not openpyxl:
        raise Exception("openpyxl no está instalado")

    safe_title = re_mod.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:50]
    timestamp = datetime.now().strftime("%Y%m%d")
    filename = f"{safe_title}_{timestamp}.xlsx"
    filepath = os.path.join(tempfile.gettempdir(), filename)

    wb = openpyxl.Workbook()
    # Eliminar la hoja default
    wb.remove(wb.active)

    # Estilos
    header_font = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill(start_color='2F5496', end_color='2F5496', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    cell_font = Font(name='Calibri', size=11)
    cell_alignment = Alignment(vertical='top', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    alt_fill = PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid')

    for sheet_info in sheets_data:
        sheet_name = sheet_info.get('sheet_name', 'Hoja1')[:31]  # Excel limit: 31 chars
        headers = sheet_info.get('headers', [])
        rows = sheet_info.get('rows', [])

        ws = wb.create_sheet(title=sheet_name)

        # Escribir encabezados
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border

        # Escribir datos
        for row_idx, row_data in enumerate(rows, 2):
            is_alt = (row_idx % 2 == 0)
            for col_idx, value in enumerate(row_data, 1):
                # Intentar convertir números
                try:
                    if isinstance(value, str):
                        if '.' in value or ',' in value:
                            value = float(value.replace(',', ''))
                        elif value.isdigit():
                            value = int(value)
                except (ValueError, AttributeError):
                    pass

                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.font = cell_font
                cell.alignment = cell_alignment
                cell.border = thin_border
                if is_alt:
                    cell.fill = alt_fill

        # Auto-ajustar anchos de columna
        for col_idx in range(1, len(headers) + 1):
            max_length = len(str(headers[col_idx - 1])) if col_idx <= len(headers) else 10
            for row in ws.iter_rows(min_row=2, min_col=col_idx, max_col=col_idx):
                for cell in row:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
            adjusted_width = min(max_length + 4, 50)
            ws.column_dimensions[get_column_letter(col_idx)].width = adjusted_width

        # Freeze header row
        ws.freeze_panes = 'A2'

        # Auto-filter
        if headers:
            ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{len(rows) + 1}"

    wb.save(filepath)
    return filepath, filename


# =====================================================
# SCHEMAS DE HERRAMIENTAS (para Anthropic API)
# =====================================================


# =====================================================
# REDDIT SCRAPING (API publica sin auth)
# =====================================================

def search_reddit(query, subreddit=None, sort='relevance', time_filter='week', limit=5):
    """Busca posts en Reddit usando la API publica JSON."""
    try:
        headers = {
            'User-Agent': 'ClaudetteBot/1.0 (personal assistant bot)',
            'Accept': 'application/json'
        }
        if subreddit:
            url = f'https://www.reddit.com/r/{subreddit}/search.json'
            params = {'q': query, 'sort': sort, 't': time_filter, 'limit': limit, 'restrict_sr': 1}
        else:
            url = 'https://www.reddit.com/search.json'
            params = {'q': query, 'sort': sort, 't': time_filter, 'limit': limit}

        resp = requests.get(url, headers=headers, params=params, timeout=15)
        resp.raise_for_status()
        data = resp.json()

        posts = data.get('data', {}).get('children', [])
        if not posts:
            return 'Sin resultados en Reddit para esa busqueda.'

        results = []
        for p in posts[:limit]:
            post = p.get('data', {})
            title = post.get('title', '')
            score = post.get('score', 0)
            comments = post.get('num_comments', 0)
            sub = post.get('subreddit', '')
            url_post = 'https://reddit.com' + post.get('permalink', '')
            selftext = post.get('selftext', '')[:300]
            item = f'\U0001f4cc **{title}**\n   r/{sub} - {score} pts - {comments} comentarios\n   {url_post}'
            if selftext and selftext.strip():
                item += f'\n   _{selftext.strip()}..._'
            results.append(item)

        return f'\U0001f534 Reddit - {len(results)} resultados para "{query}":\n\n' + '\n\n'.join(results)

    except Exception as e:
        logger.warning(f'Reddit search error: {e}')
        return f'Error buscando en Reddit: {e}'


# =====================================================
# HACKER NEWS (API oficial Firebase)
# =====================================================

def fetch_hackernews_top(limit=10, min_points=50):
    """Obtiene top stories de Hacker News con titulo, puntos y URL."""
    try:
        resp = requests.get(
            'https://hacker-news.firebaseio.com/v0/topstories.json',
            timeout=10
        )
        resp.raise_for_status()
        story_ids = resp.json()[:50]

        results = []
        for story_id in story_ids:
            if len(results) >= limit:
                break
            try:
                item_resp = requests.get(
                    f'https://hacker-news.firebaseio.com/v0/item/{story_id}.json',
                    timeout=5
                )
                item = item_resp.json()
                if not item or item.get('type') != 'story':
                    continue
                score = item.get('score', 0)
                if score < min_points:
                    continue
                title = item.get('title', '')
                url = item.get('url', f'https://news.ycombinator.com/item?id={story_id}')
                comments = item.get('descendants', 0)
                results.append(f'\U0001f536 **{title}**\n   {score} pts - {comments} comentarios\n   {url}')
            except Exception:
                continue

        if not results:
            return 'No hay stories en HN con suficientes puntos ahora.'

        return f'\U0001f536 Hacker News - Top {len(results)} stories:\n\n' + '\n\n'.join(results)

    except Exception as e:
        logger.warning(f'HN fetch error: {e}')
        return f'Error obteniendo HN: {e}'


# =====================================================
# ANALIZADOR DE CONTENIDO 8 MODOS
# =====================================================

ANALYZE_MODES_PROMPT = """Analiza en exactamente 8 modos. Directo, sin preambulos.

**1. MODELOS MENTALES**
Que marcos de pensamiento aplican (primeros principios, segundo orden, inversion, Occam, etc). Max 3 modelos aplicados al contenido especifico.

**2. DETECTOR DE HUMO**
Que no se esta diciendo. Que puede ser exagerado, falso o interesado. Que incentivos tiene quien lo publico. Sin piedad.

**3. IDEAS DE NEGOCIO**
Que oportunidades de negocio o arbitraje de informacion emergen. Al menos 2 ideas concretas y accionables.

**4. ESTRUCTURA NARRATIVA**
Como esta construido el argumento. Que tecnica persuasiva usa. Donde esta el giro o la tension.

**5. PUNTOS CIEGOS**
Que esta ignorando o dando por sentado. Que preguntas no se hace. Que perspectiva falta.

**6. PLAN DE ACCION**
Si Pablo quisiera actuar basandose en esto, que haria en los proximos 7 dias. 3 pasos concretos.

**7. SUBTEXTO**
Que se esta comunicando sin decirlo explicitamente. Tensiones, miedos, ambiciones no declaradas.

**8. CONEXION FILOSOFICA**
Conecta la idea central con un pensador o libro que Pablo conoce (Heidegger, Han, Fisher, Taleb, Jung, etc). Una conexion inesperada pero real."""


def analyze_content_deep(content, title=''):
    """
    Analiza contenido con 8 modos simultaneos de pensamiento profundo.
    Ideal para YouTube, articulos, Reddit, HN, tweets o cualquier texto.
    """
    import anthropic as _anthropic
    from config import ANTHROPIC_API_KEY, DEFAULT_MODEL
    _client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    if len(content) > 8000:
        content = content[:8000] + '\n\n[... contenido truncado]'

    title_line = f'Titulo: {title}\n' if title else ''
    prompt = f'CONTENIDO A ANALIZAR:\n{title_line}{content}\n\n---\n{ANALYZE_MODES_PROMPT}'

    try:
        response = _client.messages.create(
            model=DEFAULT_MODEL,
            max_tokens=3000,
            messages=[{'role': 'user', 'content': prompt}]
        )
        result = ''
        for block in response.content:
            if block.type == 'text':
                result += block.text
        return result if result else 'No se pudo generar el analisis.'
    except Exception as e:
        logger.error(f'analyze_content_deep error: {e}')
        return f'Error en analisis profundo: {e}'


def verify_content(url_or_text, claim=None):
    """
    Escudo de Veracidad: verifica si una noticia, URL o claim es real o fake news.
    Busca el mismo tema en 3 fuentes independientes (Reddit, HN, web) y analiza
    señales lingüísticas de desinformación. Retorna veredicto estructurado.
    """
    import anthropic as _anthropic
    from config import ANTHROPIC_API_KEY

    content_to_check = ""
    source_url = None

    # Detectar si es URL o texto libre
    stripped = url_or_text.strip()
    if stripped.startswith("http://") or stripped.startswith("https://"):
        source_url = stripped
        fetched = fetch_url(source_url)
        if fetched and not fetched.startswith("⚠️") and not fetched.startswith("❌"):
            content_to_check = fetched[:3000]
        else:
            content_to_check = claim or stripped
    else:
        content_to_check = stripped

    search_query = (claim or content_to_check)[:120]

    # Fuente 1: Reddit
    reddit_data = ""
    try:
        reddit_data = search_reddit(search_query, limit=3)[:600]
    except Exception:
        pass

    # Fuente 2: Hacker News
    hn_data = ""
    try:
        hn_data = fetch_hackernews_top(limit=5, min_points=10)[:600]
    except Exception:
        pass

    # Fuente 3: Web (Google/DuckDuckGo)
    web_data = ""
    try:
        web_data = search_web_google(search_query, max_results=3)[:600]
    except Exception:
        pass

    verification_prompt = f"""Eres un detector de fake news y desinformación experto. Analiza el siguiente contenido.

CONTENIDO A VERIFICAR:
{content_to_check[:2000]}
{("URL fuente: " + source_url) if source_url else ""}

FUENTES INDEPENDIENTES ENCONTRADAS:
--- Reddit ---
{reddit_data or "Sin resultados"}

--- Hacker News ---
{hn_data or "Sin resultados"}

--- Web ---
{web_data or "Sin resultados"}

Analiza y responde con este formato exacto:

**VEREDICTO**: [✅ VERIFICADO / ⚠️ NO CONFIRMADO / 🚨 PROBABLE FAKE / 🔍 INSUFICIENTE PARA JUZGAR]

**CONFIANZA**: X% — [una línea explicando por qué]

**SEÑALES DETECTADAS**:
- [señal 1: credibilidad o desinformación]
- [señal 2]
- [señal 3 si aplica]

**CORROBORACIÓN**:
[Qué fuentes confirman, contradicen o ignoran este claim. Si ninguna lo menciona, dilo.]

**ANÁLISIS**: [2-3 oraciones de conclusión.]

Sé directo. No inventes información que no esté en las fuentes."""

    try:
        _client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        response = _client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=1500,
            messages=[{"role": "user", "content": verification_prompt}]
        )
        result = response.content[0].text if response.content else "Sin respuesta."
        return "🛡️ *ESCUDO DE VERACIDAD*\n\n" + result
    except Exception as e:
        logger.error(f"verify_content error: {e}")
        return f"❌ Error en verificación: {e}"


TOOLS_SCHEMA = KB_TOOLS_SCHEMA + [
    {
        "name": "get_current_weather",
        "description": "Obtener el clima actual basado en latitud y longitud.",
        "input_schema": {
            "type": "object",
            "properties": {"lat": {"type": "number"}, "lon": {"type": "number"}},
            "required": ["lat", "lon"]
        }
    },
    {
        "name": "get_weather_by_city",
        "description": (
            "Obtener el clima actual por nombre de ciudad. "
            "USAR cuando Pablo mencione estar en un lugar ('estoy en Madrid', 'llegue a Tokyo') "
            "o pida el clima de una ciudad sin GPS."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "city_name": {"type": "string", "description": "Nombre de la ciudad (ej: 'San Jose, Costa Rica', 'Madrid')"}
            },
            "required": ["city_name"]
        }
    },
    {
        "name": "search_nearby_places",
        "description": "Buscar lugares cercanos (restaurantes, tiendas, etc).",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Qué buscar (ej: pizza, veterinaria)"}
            },
            "required": ["query"]
        }
    },
    {
        "name": "search_contact_and_call",
        "description": "Busca un contacto en Google y devuelve su tarjeta.",
        "input_schema": {
            "type": "object",
            "properties": {"name_query": {"type": "string"}},
            "required": ["name_query"]
        }
    },
    {
        "name": "read_book_from_drive",
        "description": "Buscar y leer documentos desde Google Drive. Soporta: Google Docs (más liviano), PDF, EPUB, TXT, MD. Usa para leer escritos de Pablo, libros, el INDICE_BIBLIOTECA, o cualquier documento en Drive.",
        "input_schema": {
            "type": "object",
            "properties": {"query": {"type": "string", "description": "Nombre o parte del nombre del libro/archivo"}},
            "required": ["query"]
        }
    },
    {
        "name": "save_user_fact",
        "description": "Guardar dato importante en memoria persistente. USA SIEMPRE esta herramienta cuando el usuario mencione: datos personales, ubicaciones, preferencias, datos de trabajo, nombres de mascotas, fechas importantes, o cuando diga 'recuerda', 'memoriza', 'anota'. Guarda proactivamente.",
        "input_schema": {
            "type": "object",
            "properties": {
                "category": {"type": "string", "description": "Categoría: Familia, Ubicacion, Preferencia, Trabajo, Personal, Fecha, etc."},
                "key": {"type": "string", "description": "Qué es (ej: 'mamá vive en', 'color favorito')"},
                "value": {"type": "string", "description": "El valor a recordar"}
            },
            "required": ["category", "key", "value"]
        }
    },
    {
        "name": "search_web",
        "description": "Buscar en Google Web.",
        "input_schema": {
            "type": "object",
            "properties": {"query": {"type": "string"}},
            "required": ["query"]
        }
    },
    {
        "name": "search_news",
        "description": "Buscar noticias recientes. Sin temas busca AI, geopolítica y mercados. Con topics personaliza.",
        "input_schema": {
            "type": "object",
            "properties": {
                "topics": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Lista de temas (opcional, ej: ['bitcoin', 'Venezuela'])"
                }
            }
        }
    },
    {
        "name": "get_calendar_events",
        "description": "Consultar calendario de Google.",
        "input_schema": {
            "type": "object",
            "properties": {
                "start_date": {"type": "string"},
                "end_date": {"type": "string"}
            },
            "required": ["start_date", "end_date"]
        }
    },
    {
        "name": "create_calendar_event",
        "description": "Agendar evento en Google Calendar.",
        "input_schema": {
            "type": "object",
            "properties": {
                "summary": {"type": "string"},
                "start_time": {"type": "string"},
                "end_time": {"type": "string"}
            },
            "required": ["summary", "start_time", "end_time"]
        }
    },
    {
        "name": "create_task",
        "description": "Crear tarea en Google Tasks.",
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {"type": "string"},
                "notes": {"type": "string"}
            },
            "required": ["title"]
        }
    },
    {
        "name": "list_tasks",
        "description": "Listar tareas de Google Tasks.",
        "input_schema": {
            "type": "object",
            "properties": {
                "show_completed": {"type": "boolean"}
            }
        }
    },
    {
        "name": "search_emails",
        "description": "Buscar correos en Gmail. Usa sintaxis Gmail: from:, to:, subject:, is:unread, has:attachment, newer_than:2d, etc.",
        "input_schema": {
            "type": "object",
            "properties": {"query": {"type": "string"}},
            "required": ["query"]
        }
    },
    {
        "name": "get_email",
        "description": "Leer el contenido completo de un email específico por su ID (obtenido de search_emails).",
        "input_schema": {
            "type": "object",
            "properties": {"email_id": {"type": "string", "description": "ID del email"}},
            "required": ["email_id"]
        }
    },
    {
        "name": "send_email",
        "description": "Enviar un email desde Gmail. Puede enviar correos nuevos o responder a existentes.",
        "input_schema": {
            "type": "object",
            "properties": {
                "to": {"type": "string", "description": "Email del destinatario"},
                "subject": {"type": "string", "description": "Asunto del correo"},
                "body": {"type": "string", "description": "Cuerpo del correo en texto plano"},
                "reply_to_id": {"type": "string", "description": "ID del email al que responder (opcional)"}
            },
            "required": ["to", "subject", "body"]
        }
    },
    {
        "name": "generate_image",
        "description": "Generar una imagen usando AI (DALL-E 3) basada en una descripción.",
        "input_schema": {
            "type": "object",
            "properties": {"prompt": {"type": "string", "description": "Descripción visual detallada"}},
            "required": ["prompt"]
        }
    },
    {
        "name": "read_local_file",
        "description": "Leer archivos del sistema de modelos mentales. Archivos disponibles: MODELS_DEEP.md (176 modelos adicionales por dominio), FRAMEWORK.md (metodología paso-a-paso), ANTIPATTERNS.md (cuándo NO usar modelos), TEMPLATES.md (plantillas ejecutables para decisiones, negocios, riesgo, ética, innovación). USAR cuando análisis profundo requiere más de los 40 modelos core.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Nombre del archivo: MODELS_DEEP.md, FRAMEWORK.md, ANTIPATTERNS.md, o TEMPLATES.md",
                    "enum": ["MODELS_DEEP.md", "FRAMEWORK.md", "ANTIPATTERNS.md", "TEMPLATES.md"]
                }
            },
            "required": ["filename"]
        }
    },
    {
        "name": "fetch_url",
        "description": "Lee el contenido de una URL/link que Pablo comparta. Funciona con artículos web, tweets/posts de X (Twitter), blogs, noticias, y cualquier página pública. Usa cuando Pablo mande un link y quiera que lo leas o analices.",
        "input_schema": {
            "type": "object",
            "properties": {
                "url": {
                    "type": "string",
                    "description": "URL completa a leer (con https://)"
                }
            },
            "required": ["url"]
        }
    },
    {
        "name": "generate_document",
        "description": """Genera un documento largo y descargable (.docx o .md) que se envía como archivo adjunto en Telegram. 
USA ESTA HERRAMIENTA cuando Pablo pida:
- Generar un reporte, informe, ensayo, bitácora, resumen extenso, documento de trabajo
- Compilar conversaciones o impresiones en un documento
- Cualquier texto que supere las limitaciones de un mensaje de Telegram (~4000 chars)
- 'Hazme un documento', 'generame un reporte', 'arma la bitácora', 'compila esto en un doc'

IMPORTANTE: El campo 'content' es el documento COMPLETO que quieres generar. Usa formato Markdown:
- # Título, ## Sección, ### Subsección
- **negrita**, *cursiva*
- Listas con - o 1. 2. 3.
- > para citas
- --- para separadores

El contenido se convierte automáticamente en un Word (.docx) con formato profesional.""",
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Título del documento"
                },
                "content": {
                    "type": "string",
                    "description": "Contenido completo del documento en formato Markdown. Puede ser tan largo como sea necesario."
                },
                "format": {
                    "type": "string",
                    "description": "Formato del archivo: 'docx' (Word, default) o 'md' (Markdown)",
                    "enum": ["docx", "md"],
                    "default": "docx"
                }
            },
            "required": ["title", "content"]
        }
    },
    {
        "name": "generate_spreadsheet",
        "description": """Genera un archivo Excel (.xlsx) con formato profesional que se envía como archivo adjunto en Telegram.
USA ESTA HERRAMIENTA cuando Pablo pida:
- Tablas, comparativas, matrices de datos
- Hojas de cálculo, presupuestos, tracking, inventarios
- Análisis comparativo en formato tabular
- 'Hazme una tabla en Excel', 'ponlo en una hoja de cálculo', 'generame un spreadsheet'

Soporta múltiples hojas en un solo archivo. El Excel se genera con:
- Encabezados azules con texto blanco
- Filas alternadas en gris claro
- Auto-filtros en cada columna
- Fila de encabezado congelada
- Anchos de columna auto-ajustados""",
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Título del archivo Excel"
                },
                "sheets": {
                    "type": "array",
                    "description": "Lista de hojas del Excel. Cada hoja tiene nombre, encabezados y filas de datos.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "sheet_name": {
                                "type": "string",
                                "description": "Nombre de la hoja (max 31 caracteres)"
                            },
                            "headers": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Lista de encabezados de columna"
                            },
                            "rows": {
                                "type": "array",
                                "items": {
                                    "type": "array",
                                    "items": {"type": "string"}
                                },
                                "description": "Lista de filas, cada fila es una lista de valores (strings)"
                            }
                        },
                        "required": ["sheet_name", "headers", "rows"]
                    }
                }
            },
            "required": ["title", "sheets"]
        }
    },
    {
        "name": "search_library",
        "description": "Buscar en la biblioteca personal de Pablo (2100+ libros). Busca por tema, concepto, autor o tag. Usa para responder preguntas sobre filosofía, ciencias sociales, esoterismo, psicología, y cualquier tema de sus libros. Retorna extractos relevantes con contexto.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Tema, concepto, autor o palabra clave a buscar"},
                "limit": {"type": "integer", "description": "Cantidad de resultados (default 5)", "default": 5}
            },
            "required": ["query"]
        }
    },
    {
        "name": "search_library_by_author",
        "description": "Listar todos los libros de un autor específico en la biblioteca de Pablo.",
        "input_schema": {
            "type": "object",
            "properties": {
                "author": {"type": "string", "description": "Nombre del autor"}
            },
            "required": ["author"]
        }
    },
    {
        "name": "search_library_by_tag",
        "description": "Buscar libros por tag/etiqueta en la biblioteca. Tags incluyen: filosofia, psicologia, nihilismo, democracia, capitalismo, atencion, zen, contemplacion, etc.",
        "input_schema": {
            "type": "object",
            "properties": {
                "tag": {"type": "string", "description": "Tag a buscar (minúsculas, sin espacios)"}
            },
            "required": ["tag"]
        }
    },
    {
        "name": "get_book_detail",
        "description": "Obtener el extracto completo de un libro específico de la biblioteca. Usa cuando Pablo pide profundizar en un libro encontrado.",
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Título o parte del título del libro"}
            },
            "required": ["title"]
        }
    },
    {
        "name": "library_stats",
        "description": "Mostrar estadísticas de la biblioteca de Pablo (total libros, autores, categorías).",
        "input_schema": {
            "type": "object",
            "properties": {}
        }
    },
    {
        "name": "search_reddit",
        "description": "Busca posts en Reddit sobre cualquier tema. Util para saber debates actuales, noticias de nicho, opiniones sobre tecnologia, trading, filosofia, etc.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Termino de busqueda"},
                "subreddit": {"type": "string", "description": "Subreddit especifico (opcional)"},
                "sort": {"type": "string", "description": "Orden: relevance, new, top, hot"},
                "time_filter": {"type": "string", "description": "Filtro: day, week, month, year, all"},
                "limit": {"type": "integer", "description": "Cantidad de resultados (default: 5)"}
            },
            "required": ["query"]
        }
    },
    {
        "name": "fetch_hackernews_top",
        "description": "Obtiene top stories actuales de Hacker News. Mejor fuente para noticias de tecnologia, startups, IA, ciencias y cultura tech.",
        "input_schema": {
            "type": "object",
            "properties": {
                "limit": {"type": "integer", "description": "Cantidad de stories (default: 10)"},
                "min_points": {"type": "integer", "description": "Minimo de puntos (default: 50)"}
            }
        }
    },
    {
        "name": "analyze_content_deep",
        "description": "Analiza cualquier contenido con 8 modos simultaneos: modelos mentales, detector de humo, ideas de negocio, estructura narrativa, puntos ciegos, plan de accion, subtexto, conexion filosofica. Usar cuando Pablo diga analiza esto, 8 modos, analisis profundo, o mande contenido extenso.",
        "input_schema": {
            "type": "object",
            "properties": {
                "content": {"type": "string", "description": "Texto o contenido a analizar"},
                "title": {"type": "string", "description": "Titulo del contenido (opcional)"}
            },
            "required": ["content"]
        }
    },
    {
        "name": "verify_content",
        "description": (
            "Escudo de Veracidad: verifica si una noticia, URL o claim es real o fake news. "
            "Busca el mismo tema en 3 fuentes independientes (Reddit, HN, web) y analiza señales de desinformación. "
            "ACTIVAR AUTOMATICAMENTE cuando Pablo comparta: una URL de noticias, un claim que suene dudoso, "
            "una noticia viral, algo de X/Twitter, o cuando diga 'verifica esto', 'es fake?', 'es verdad?', "
            "'chequea esto'. Retorna veredicto: VERIFICADO / NO CONFIRMADO / PROBABLE FAKE."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "url_or_text": {
                    "type": "string",
                    "description": "URL de la noticia a verificar, o texto/claim directo"
                },
                "claim": {
                    "type": "string",
                    "description": "El claim específico a verificar (opcional, se extrae del contenido si no se da)"
                }
            },
            "required": ["url_or_text"]
        }
    }
]


# =====================================================
# EJECUCIÃ“N DE HERRAMIENTAS
# =====================================================

# Estado global compartido con main.py
user_locations = {}


async def execute_tool(tool_name: str, tool_input: dict, chat_id: int, context):
    """Ejecuta una herramienta y retorna el resultado como string."""

    try:
        if tool_name == "get_current_weather":
            return get_weather(tool_input['lat'], tool_input['lon'])

        elif tool_name == "get_weather_by_city":
            return get_weather_by_city(tool_input["city_name"])

        elif tool_name == "search_contact_and_call":
            query = tool_input['name_query']
            results = google_contacts.search_contact(query)
            if not results:
                return f"âŒ No encontré a '{query}'."
            contact = results[0]
            await context.bot.send_contact(
                chat_id=chat_id,
                phone_number=contact['phone'],
                first_name=contact['name']
            )
            return f"✅ Contacto: {contact['name']}."

        elif tool_name == "search_nearby_places":
            loc = user_locations.get(chat_id, DEFAULT_LOCATION)
            lat, lng = loc['lat'], loc['lng']
            loc_name = loc.get('name', 'San José, Costa Rica')
            query = tool_input['query']

            try:
                places_result = google_places.search_nearby_places(query, lat, lng)
                if not places_result or "⚠️" in str(places_result):
                    raise Exception(str(places_result))
                return places_result
            except Exception as e:
                logger.warning(f"Places API falló: {e}. Fallback a web search.")
                fallback_query = f"{query} near {loc_name}"
                web_result = search_web_google(fallback_query)
                if web_result and "no devolvió resultados" not in web_result:
                    return f"ðŸ” (Búsqueda web, Places API no disponible):\n{web_result}"
                return f"ðŸ” (Búsqueda general):\n{search_web_google(f'{query} Costa Rica')}"

        elif tool_name == "read_book_from_drive":
            return read_book_from_drive(tool_input['query'])

        elif tool_name == "save_user_fact":
            full_key = f"{tool_input.get('category', 'General')}: {tool_input.get('key', 'Dato')}"
            save_fact(full_key, tool_input['value'])
            return f"✅ Guardado: {full_key}"

        elif tool_name == "search_web":
            return search_web_google(tool_input['query'])

        elif tool_name == "search_news":
            topics = tool_input.get('topics')
            return search_news(topics)

        elif tool_name == "get_calendar_events":
            return google_calendar.get_calendar_events(
                clean_date_iso(tool_input['start_date']),
                clean_date_iso(tool_input['end_date'], True)
            )

        elif tool_name == "create_calendar_event":
            return google_calendar.create_calendar_event(
                tool_input['summary'],
                clean_date_iso(tool_input['start_time']),
                clean_date_iso(tool_input['end_time'])
            )

        elif tool_name == "create_task":
            return google_tasks.create_task(
                tool_input['title'],
                tool_input.get('notes')
            )

        elif tool_name == "list_tasks":
            return google_tasks.list_tasks(tool_input.get('show_completed', False))

        elif tool_name == "search_emails":
            return gmail_service.search_emails(tool_input['query'])

        elif tool_name == "get_email":
            return gmail_service.get_email(tool_input['email_id'])

        elif tool_name == "send_email":
            return gmail_service.send_email(
                to=tool_input['to'],
                subject=tool_input['subject'],
                body=tool_input['body'],
                reply_to_id=tool_input.get('reply_to_id')
            )

        elif tool_name == "generate_image":
            if not openai_client:
                return "OpenAI no configurado."
            msg = await context.bot.send_message(chat_id, "ðŸŽ¨ Pintando tu idea...")
            response = openai_client.images.generate(
                model="dall-e-3",
                prompt=tool_input['prompt'],
                size="1024x1024",
                quality="standard",
                n=1
            )
            await context.bot.send_photo(
                chat_id,
                photo=response.data[0].url,
                caption=f"ðŸŽ¨ {tool_input['prompt']}"
            )
            await context.bot.delete_message(chat_id, msg.message_id)
            return "✅ Imagen generada y enviada."

        elif tool_name == "read_local_file":
            return read_local_file(tool_input['filename'])

        elif tool_name == "fetch_url":
            return fetch_url(tool_input['url'])

        elif tool_name == "generate_document":
            doc_format = tool_input.get('format', 'docx')
            title = tool_input['title']
            content = tool_input['content']

            msg = await context.bot.send_message(chat_id, "ðŸ“ Generando documento...")

            try:
                filepath, filename = generate_document(title, content, doc_format)

                # Enviar como archivo adjunto en Telegram
                with open(filepath, 'rb') as f:
                    await context.bot.send_document(
                        chat_id=chat_id,
                        document=f,
                        filename=filename,
                        caption=f"ðŸ“„ {title}"
                    )

                # Limpiar archivo temporal
                os.unlink(filepath)
                await context.bot.delete_message(chat_id, msg.message_id)

                return f"✅ Documento '{title}' generado y enviado como {filename}"

            except Exception as e:
                logger.error(f"Document generation error: {e}")
                await context.bot.delete_message(chat_id, msg.message_id)
                return f"⚠️ Error generando documento: {e}"

        elif tool_name == "generate_spreadsheet":
            title = tool_input['title']
            sheets_data = tool_input['sheets']

            msg = await context.bot.send_message(chat_id, "ðŸ“Š Generando Excel...")

            try:
                filepath, filename = generate_spreadsheet(title, sheets_data)

                with open(filepath, 'rb') as f:
                    await context.bot.send_document(
                        chat_id=chat_id,
                        document=f,
                        filename=filename,
                        caption=f"ðŸ“Š {title}"
                    )

                os.unlink(filepath)
                await context.bot.delete_message(chat_id, msg.message_id)

                total_rows = sum(len(s.get('rows', [])) for s in sheets_data)
                total_sheets = len(sheets_data)
                return f"✅ Excel '{title}' generado: {total_sheets} hoja(s), {total_rows} filas"

            except Exception as e:
                logger.error(f"Spreadsheet generation error: {e}")
                await context.bot.delete_message(chat_id, msg.message_id)
                return f"⚠️ Error generando Excel: {e}"

        elif tool_name == "search_library":
            return search_library(tool_input['query'], tool_input.get('limit', 5))

        elif tool_name == "search_library_by_author":
            return search_by_author(tool_input['author'])

        elif tool_name == "search_library_by_tag":
            return search_by_tag(tool_input['tag'])

        elif tool_name == "get_book_detail":
            return get_book_content(tool_input['title'])

        elif tool_name == "library_stats":
            return get_library_stats()

        elif tool_name == "search_reddit":
            return search_reddit(
                tool_input["query"],
                subreddit=tool_input.get("subreddit"),
                sort=tool_input.get("sort", "relevance"),
                time_filter=tool_input.get("time_filter", "week"),
                limit=tool_input.get("limit", 5)
            )

        elif tool_name == "fetch_hackernews_top":
            return fetch_hackernews_top(
                limit=tool_input.get("limit", 10),
                min_points=tool_input.get("min_points", 50)
            )

        elif tool_name == "analyze_content_deep":
            return analyze_content_deep(
                tool_input["content"],
                title=tool_input.get("title", "")
            )

        elif tool_name == "verify_content":
            return verify_content(
                tool_input["url_or_text"],
                claim=tool_input.get("claim")
            )

        # Knowledge Base tools
        if tool_name.startswith('kb_') or tool_name == "search_everything":
            return await execute_kb_tool(tool_name, tool_input)

        return f"Herramienta '{tool_name}' no encontrada."

    except Exception as e:
        logger.error(f"Tool Error {tool_name}: {e}")
        return f"Error en herramienta {tool_name}: {e}"



